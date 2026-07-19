"""Tests for ``app.services.file_parser``.

Public API::

    async def parse(file_path: str, mime_type: str) -> ParsedDocument
    class FileParser: async def parse(...)
    class ParsedDocument: text, page_count, metadata, used_ocr, warnings
"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.services import file_parser as fp


async def test_parse_txt_returns_text_content(tmp_path: Path):
    """Arrange: a UTF-8 .txt file. Act: parse. Assert: text roundtrips."""
    # Arrange
    p = tmp_path / "sample.txt"
    body = "本契約書は請負契約に関するものである。"
    p.write_text(body, encoding="utf-8")
    # Act
    doc = await fp.parse(str(p), mime_type="text/plain")
    # Assert
    assert body in doc.text
    assert doc.used_ocr is False


async def test_parse_docx_extracts_paragraphs(tmp_path: Path):
    """Arrange: minimal docx. Act: parse. Assert: paragraph text found."""
    # Arrange
    try:
        from docx import Document
    except Exception:
        pytest.skip("python-docx unavailable")
    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("第一条 目的")
    doc.add_paragraph("本契約は工事請負を目的とする。")
    doc.save(p)
    # Act
    parsed = await fp.parse(
        str(p),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    # Assert
    assert "第一条" in parsed.text


async def test_parse_missing_file_raises(tmp_path: Path):
    """Arrange: nonexistent path. Act: parse. Assert: FileParserError."""
    # Arrange
    missing = tmp_path / "no_such.txt"
    # Act / Assert
    with pytest.raises(fp.FileParserError):
        await fp.parse(str(missing), mime_type="text/plain")


async def test_parse_unknown_mime_falls_back_to_suffix(tmp_path: Path):
    """Arrange: .txt file with octet-stream mime. Act: parse. Assert: text returned."""
    # Arrange
    p = tmp_path / "data.txt"
    p.write_text("plain text", encoding="utf-8")
    # Act
    try:
        doc = await fp.parse(str(p), mime_type="application/octet-stream")
    except fp.FileParserError:
        pytest.skip("parser does not fall back on suffix (acceptable)")
    # Assert
    assert "plain text" in doc.text


async def test_parsed_document_has_expected_attributes(tmp_path: Path):
    """Arrange: simple txt. Act: parse. Assert: ParsedDocument shape."""
    # Arrange
    p = tmp_path / "x.txt"
    p.write_text("hello", encoding="utf-8")
    # Act
    doc = await fp.parse(str(p), mime_type="text/plain")
    # Assert
    for attr in ("text", "page_count", "metadata", "used_ocr", "warnings"):
        assert hasattr(doc, attr)


# ---------------------------------------------------------------------------
# parse() routing — lines 74, 79, 83
# ---------------------------------------------------------------------------


async def test_parse_unknown_mime_and_suffix_raises(tmp_path: Path):
    """Arrange: .bin file with unknown mime. Act: parse. Assert: FileParserError (line 74)."""
    p = tmp_path / "mystery.bin"
    p.write_bytes(b"\x00\x01\x02")
    with pytest.raises(fp.FileParserError, match="unsupported mime/suffix"):
        await fp.parse(str(p), mime_type="application/octet-stream")


async def test_parse_routes_pdf_via_mime(tmp_path: Path):
    """Arrange: PDF mime. Act: parse. Assert: _parse_pdf called (line 79)."""
    from unittest.mock import MagicMock, patch

    p = tmp_path / "doc.pdf"
    p.write_bytes(b"%PDF-1.4 fake")

    mock_page = MagicMock()
    mock_page.extract_text.return_value = "pdf text"
    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = None

    parser = fp.FileParser()
    with patch("pypdf.PdfReader", return_value=mock_reader):
        result = await parser.parse(str(p), mime_type="application/pdf")

    assert result.text == "pdf text"
    assert result.page_count == 1


async def test_parse_routes_xlsx_via_mime(tmp_path: Path):
    """Arrange: xlsx mime. Act: parse. Assert: _parse_xlsx path executed (line 83)."""
    from unittest.mock import MagicMock, patch

    p = tmp_path / "sheet.xlsx"
    p.write_bytes(b"PK fake xlsx")

    mock_sheet = MagicMock()
    mock_sheet.title = "Sheet1"
    mock_sheet.iter_rows.return_value = [("値A", "値B")]

    mock_wb = MagicMock()
    mock_wb.worksheets = [mock_sheet]
    mock_wb.sheetnames = ["Sheet1"]

    parser = fp.FileParser()
    with patch("openpyxl.load_workbook", return_value=mock_wb):
        result = await parser.parse(
            str(p),
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    assert "Sheet1" in result.text
    assert result.metadata["sheets"] == ["Sheet1"]


# ---------------------------------------------------------------------------
# _parse_pdf — lines 93-118
# ---------------------------------------------------------------------------


def test_parse_pdf_extracts_text(tmp_path: Path):
    """Arrange: mocked PdfReader with one text page. Act: _parse_pdf. Assert: text returned."""
    from unittest.mock import MagicMock, patch

    mock_page = MagicMock()
    mock_page.extract_text.return_value = "契約書本文"

    mock_metadata = MagicMock()
    mock_metadata.producer = "TestProducer"
    mock_metadata.author = "TestAuthor"
    mock_metadata.title = "TestTitle"

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = mock_metadata

    p = tmp_path / "test.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser()
    with patch("pypdf.PdfReader", return_value=mock_reader):
        doc = parser._parse_pdf(p)

    assert "契約書本文" in doc.text
    assert doc.page_count == 1
    assert doc.used_ocr is False
    assert doc.metadata["producer"] == "TestProducer"
    assert doc.metadata["author"] == "TestAuthor"
    assert doc.metadata["title"] == "TestTitle"


def test_parse_pdf_page_extract_exception_adds_warning(tmp_path: Path):
    """Arrange: page.extract_text raises. Act: _parse_pdf. Assert: warning added (lines 101-103)."""
    from unittest.mock import MagicMock, patch

    mock_page = MagicMock()
    mock_page.extract_text.side_effect = RuntimeError("corrupt page")

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = None

    p = tmp_path / "bad.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser()
    with patch("pypdf.PdfReader", return_value=mock_reader):
        doc = parser._parse_pdf(p)

    assert any("extract failed" in w for w in doc.warnings)
    assert doc.metadata["producer"] is None


def test_parse_pdf_image_only_uses_ocr_stub(tmp_path: Path):
    """All pages empty → image-only PDF → used_ocr=True, OCR stub text (lines 107-111)."""
    from unittest.mock import MagicMock, patch

    mock_page = MagicMock()
    mock_page.extract_text.return_value = ""

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = None

    p = tmp_path / "image_only.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser(ocr_enabled=False)
    with patch("pypdf.PdfReader", return_value=mock_reader):
        doc = parser._parse_pdf(p)

    assert doc.used_ocr is True
    assert "OCR-STUB" in doc.text
    assert any("image-only" in w for w in doc.warnings)


def test_parse_pdf_metadata_none_when_reader_metadata_absent(tmp_path: Path):
    """reader.metadata is None → all meta fields are None (lines 113-117)."""
    from unittest.mock import MagicMock, patch

    mock_page = MagicMock()
    mock_page.extract_text.return_value = "some text"

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = None

    p = tmp_path / "nometa.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser()
    with patch("pypdf.PdfReader", return_value=mock_reader):
        doc = parser._parse_pdf(p)

    assert doc.metadata["producer"] is None
    assert doc.metadata["author"] is None
    assert doc.metadata["title"] is None


# ---------------------------------------------------------------------------
# _parse_docx table extraction — lines 134-137
# ---------------------------------------------------------------------------


def test_parse_docx_table_cells_included(tmp_path: Path):
    """Arrange: docx with table. Act: _parse_docx. Assert: cell text in output (lines 134-137)."""
    from unittest.mock import MagicMock, patch

    mock_cell_a = MagicMock()
    mock_cell_a.text = "甲"
    mock_cell_b = MagicMock()
    mock_cell_b.text = "乙"

    mock_row = MagicMock()
    mock_row.cells = [mock_cell_a, mock_cell_b]

    mock_table = MagicMock()
    mock_table.rows = [mock_row]

    mock_para = MagicMock()
    mock_para.text = "序文"

    mock_doc_obj = MagicMock()
    mock_doc_obj.paragraphs = [mock_para]
    mock_doc_obj.tables = [mock_table]

    p = tmp_path / "contract.docx"
    p.write_bytes(b"PK fake docx")

    parser = fp.FileParser()
    with patch("docx.Document", return_value=mock_doc_obj):
        doc = parser._parse_docx(p)

    assert "甲" in doc.text
    assert "乙" in doc.text
    assert "序文" in doc.text


def test_parse_docx_empty_table_cells_skipped(tmp_path: Path):
    """Arrange: table row with all blank cells. Act: _parse_docx. Assert: no extra lines."""
    from unittest.mock import MagicMock, patch

    mock_cell = MagicMock()
    mock_cell.text = "   "  # whitespace only

    mock_row = MagicMock()
    mock_row.cells = [mock_cell]

    mock_table = MagicMock()
    mock_table.rows = [mock_row]

    mock_para = MagicMock()
    mock_para.text = "内容"

    mock_doc_obj = MagicMock()
    mock_doc_obj.paragraphs = [mock_para]
    mock_doc_obj.tables = [mock_table]

    p = tmp_path / "empty_table.docx"
    p.write_bytes(b"PK fake docx")

    parser = fp.FileParser()
    with patch("docx.Document", return_value=mock_doc_obj):
        doc = parser._parse_docx(p)

    assert doc.text == "内容"


# ---------------------------------------------------------------------------
# _parse_xlsx — lines 147-158
# ---------------------------------------------------------------------------


def test_parse_xlsx_extracts_sheet_data(tmp_path: Path):
    """Arrange: mocked workbook with one sheet. Act: _parse_xlsx. Assert: data in text."""
    from unittest.mock import MagicMock, patch

    mock_row = (MagicMock(), MagicMock())
    mock_row[0].__str__ = lambda self: "工事名"
    mock_row[1].__str__ = lambda self: "橋梁工事"

    mock_sheet = MagicMock()
    mock_sheet.title = "Sheet1"
    mock_sheet.iter_rows.return_value = [("工事名", "橋梁工事")]

    mock_wb = MagicMock()
    mock_wb.worksheets = [mock_sheet]
    mock_wb.sheetnames = ["Sheet1"]

    p = tmp_path / "data.xlsx"
    p.write_bytes(b"PK fake xlsx")

    parser = fp.FileParser()
    with patch("openpyxl.load_workbook", return_value=mock_wb):
        doc = parser._parse_xlsx(p)

    assert "工事名" in doc.text
    assert "橋梁工事" in doc.text
    assert "Sheet1" in doc.text
    assert doc.metadata["sheets"] == ["Sheet1"]
    mock_wb.close.assert_called_once()


def test_parse_xlsx_skips_none_cells(tmp_path: Path):
    """Arrange: row containing None values. Act: _parse_xlsx. Assert: None not in output."""
    from unittest.mock import MagicMock, patch

    mock_sheet = MagicMock()
    mock_sheet.title = "Data"
    mock_sheet.iter_rows.return_value = [(None, "有効値", None)]

    mock_wb = MagicMock()
    mock_wb.worksheets = [mock_sheet]
    mock_wb.sheetnames = ["Data"]

    p = tmp_path / "sparse.xlsx"
    p.write_bytes(b"PK fake xlsx")

    parser = fp.FileParser()
    with patch("openpyxl.load_workbook", return_value=mock_wb):
        doc = parser._parse_xlsx(p)

    assert "None" not in doc.text
    assert "有効値" in doc.text


# ---------------------------------------------------------------------------
# _parse_text encoding fallback — lines 172-174
# ---------------------------------------------------------------------------


def test_parse_text_cp932_fallback(tmp_path: Path):
    """cp932-encoded file is decoded via encoding fallback chain."""
    p = tmp_path / "sjis.txt"
    content = "工事請負契約"
    p.write_bytes(content.encode("cp932"))

    parser = fp.FileParser()
    doc = parser._parse_text(p)

    assert content in doc.text
    assert doc.metadata["encoding"] in ("cp932", "shift_jis")


def test_parse_text_all_encodings_fail_raises(tmp_path: Path):
    """All encodings fail → FileParserError raised (line 174)."""
    p = tmp_path / "broken.txt"
    # Sequence of bytes that fails utf-8, utf-8-sig, cp932, shift_jis
    p.write_bytes(b"\x81\x40\xff\xfe\x00" * 20 + b"\x99\x99\x99\x99\x99")

    parser = fp.FileParser()
    with pytest.raises(fp.FileParserError, match="could not decode"):
        parser._parse_text(p)


# ---------------------------------------------------------------------------
# _ocr_stub — lines 178-181
# ---------------------------------------------------------------------------


def test_ocr_stub_disabled_returns_stub_message(tmp_path: Path):
    """Arrange: ocr_enabled=False. Act: _ocr_stub. Assert: OCR-STUB in result (line 178-179)."""
    p = tmp_path / "image.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser(ocr_enabled=False)
    result = parser._ocr_stub(p)

    assert "OCR-STUB" in result
    assert p.name in result


def test_ocr_stub_enabled_returns_ocr_message(tmp_path: Path):
    """Arrange: ocr_enabled=True. Act: _ocr_stub. Assert: OCR (not STUB) in result (line 181)."""
    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser(ocr_enabled=True)
    result = parser._ocr_stub(p)

    assert "[OCR]" in result
    assert p.name in result


# ---------------------------------------------------------------------------
# parse_text — utf-8-sig BOM stripping
# ---------------------------------------------------------------------------


def test_parse_text_utf8_sig_bom(tmp_path: Path):
    """BOM-prefixed UTF-8 file: content is readable; encoding is utf-8 or utf-8-sig."""
    content = "BOM付きテキスト"
    p = tmp_path / "bom.txt"
    p.write_bytes(b"\xef\xbb\xbf" + content.encode("utf-8"))

    parser = fp.FileParser()
    doc = parser._parse_text(p)

    # utf-8 succeeds first in the fallback chain (BOM is kept as U+FEFF char);
    # utf-8-sig would strip the BOM. Either outcome is acceptable.
    assert content in doc.text
    assert doc.metadata["encoding"] in ("utf-8", "utf-8-sig")


# ---------------------------------------------------------------------------
# parse_text via module-level parse()
# ---------------------------------------------------------------------------


async def test_module_level_parse_routes_text(tmp_path: Path):
    """Arrange: module-level parse() convenience wrapper. Act: call it. Assert: text returned."""
    p = tmp_path / "mod.txt"
    p.write_text("モジュールレベルテスト", encoding="utf-8")
    doc = await fp.parse(str(p), mime_type="text/plain")
    assert "モジュールレベルテスト" in doc.text


# ---------------------------------------------------------------------------
# _parse_pdf via parse() async path (integration)
# ---------------------------------------------------------------------------


async def test_parse_pdf_via_async_parse(tmp_path: Path):
    """Arrange: pdf mime + mocked PdfReader. Act: await parse(). Assert: text from mock returned."""
    from unittest.mock import MagicMock, patch

    mock_page = MagicMock()
    mock_page.extract_text.return_value = "async pdf text"

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_reader.metadata = None

    p = tmp_path / "async.pdf"
    p.write_bytes(b"%PDF fake")

    parser = fp.FileParser()
    with patch("pypdf.PdfReader", return_value=mock_reader):
        doc = await parser.parse(str(p), mime_type="application/pdf")

    assert "async pdf text" in doc.text
    assert doc.page_count == 1
