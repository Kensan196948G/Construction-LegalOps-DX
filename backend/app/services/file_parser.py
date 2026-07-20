"""File-content parser.

Extracts plain text from PDF, Word (``.docx``), and Excel (``.xlsx``)
uploads so downstream services (AI review, clause extractor) can operate
on a normalised string. Image-only PDFs fail closed until a real OCR backend
is explicitly approved and configured; placeholder OCR text must never flow
into legal review.

The parser is async at the API boundary (callers may stream large files
from object storage) but the heavy lifting itself is synchronous — we
run it on the default thread pool to avoid blocking the event loop.
"""

from __future__ import annotations

import asyncio
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Final


class FileParserError(RuntimeError):
    """Raised when a file cannot be parsed."""


@dataclass(slots=True)
class ParsedDocument:
    """Structured output of :meth:`FileParser.parse`."""

    text: str
    page_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)
    used_ocr: bool = False
    warnings: list[str] = field(default_factory=list)


# MIME type → backend name mapping.
_MIME_TO_BACKEND: Final[dict[str, str]] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",  # legacy .doc — best effort, may warn
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/plain": "text",
    "text/csv": "text",
}

# Suffix fallback when MIME is missing/unknown.
_SUFFIX_TO_BACKEND: Final[dict[str, str]] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".txt": "text",
    ".csv": "text",
}


@dataclass(slots=True)
class FileParser:
    """Pluggable file parser."""

    ocr_enabled: bool = False  # Loop 4 will flip this to True

    async def parse(self, file_path: str, mime_type: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileParserError(f"file not found: {file_path}")

        backend = _MIME_TO_BACKEND.get(mime_type) or _SUFFIX_TO_BACKEND.get(
            path.suffix.lower()
        )
        if backend is None:
            raise FileParserError(
                f"unsupported mime/suffix: mime={mime_type}, suffix={path.suffix}"
            )

        if backend == "pdf":
            return await asyncio.to_thread(self._parse_pdf, path)
        if backend == "docx":
            return await asyncio.to_thread(self._parse_docx, path)
        if backend == "xlsx":
            return await asyncio.to_thread(self._parse_xlsx, path)
        if backend == "text":
            return await asyncio.to_thread(self._parse_text, path)
        raise FileParserError(f"unhandled backend: {backend}")  # pragma: no cover

    # ------------------------------------------------------------------
    # Backend implementations
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader  # local import keeps top-level fast

        reader = PdfReader(str(path))
        pages_text: list[str] = []
        warnings: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                pages_text.append(page.extract_text() or "")
            except Exception as exc:
                warnings.append(f"page {i}: extract failed ({exc!s})")
                pages_text.append("")

        body = "\n\n".join(p for p in pages_text if p.strip())
        if not body.strip():
            # No extractable text — image-only PDF. Returning placeholder OCR
            # text would make downstream legal review look evidence-backed, so
            # fail closed until a real OCR provider is configured.
            warnings.append("PDF appears image-only; OCR backend is not configured")
            self._raise_ocr_unavailable(path, warnings=warnings)

        meta = {
            "producer": getattr(reader.metadata, "producer", None) if reader.metadata else None,
            "author": getattr(reader.metadata, "author", None) if reader.metadata else None,
            "title": getattr(reader.metadata, "title", None) if reader.metadata else None,
        }
        return ParsedDocument(
            text=body,
            page_count=len(reader.pages),
            metadata=meta,
            used_ocr=False,
            warnings=warnings,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Include simple table contents because contracts often place
        # 当事者表示 in a header table.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n".join(paragraphs)
        return ParsedDocument(
            text=text,
            page_count=0,  # docx is reflowable; pages are not stable
            metadata={"paragraphs": len(doc.paragraphs)},
        )

    def _parse_xlsx(self, path: Path) -> ParsedDocument:
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True, read_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            lines.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append("\t".join(cells))
        wb.close()
        return ParsedDocument(
            text="\n".join(lines),
            page_count=len(wb.sheetnames),
            metadata={"sheets": wb.sheetnames},
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        encodings = ("utf-8", "utf-8-sig", "cp932", "shift_jis")
        for enc in encodings:
            try:
                return ParsedDocument(
                    text=path.read_text(encoding=enc),
                    metadata={"encoding": enc},
                )
            except UnicodeDecodeError:
                continue
        raise FileParserError(f"could not decode text file: {path}")

    def _raise_ocr_unavailable(self, path: Path, *, warnings: list[str] | None = None) -> None:
        """Fail closed for image-only PDFs until real OCR is wired."""
        detail = "; ".join(warnings or [])
        suffix = f" ({detail})" if detail else ""
        raise FileParserError(
            f"OCR backend is not configured for image-only PDF: {path.name}{suffix}"
        )


# Module-level convenience wrapper.
_default_parser = FileParser()


async def parse(file_path: str, mime_type: str) -> ParsedDocument:
    """Parse a file using the default :class:`FileParser`."""
    return await _default_parser.parse(file_path, mime_type)
